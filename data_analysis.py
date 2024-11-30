#Import Required Libraries

import pandas as pd
import warnings
warnings.filterwarnings("ignore")
import time
import os
import docx
from docx.shared import Inches
from datetime import datetime

#importing folder and file functions from Folder_file_actions.folder_file
from Features.Folder_file_actions.folder_file import create_folder
from Features.Folder_file_actions.folder_file import select_file
#from Features.Folder_file_actions.folder_file import process_file_in_folder

#importing dataset functions from Text_analysis.dataset_functions
from Features.Text_analysis.dataset_functions import format_column_names
from Features.Text_analysis.dataset_functions import calculate_filled_fields
from Features.Text_analysis.dataset_functions import missing_values_table
from Features.Text_analysis.dataset_functions import handle_missing_values
from Features.Text_analysis.dataset_functions import select_columns_for_analysis
from Features.Text_analysis.dataset_functions import get_top_values

#importing text functions from Text_analysis.text_functions
from Features.Text_analysis.text_functions import get_common_words
from Features.Text_analysis.text_functions import stem_sentence
from Features.Text_analysis.text_functions import remove_freqwords
from Features.Text_analysis.text_functions import remove_rarewords
from Features.Text_analysis.text_functions import remove_calendar
from Features.Text_analysis.text_functions import remove_greetings
from Features.Text_analysis.text_functions import process_text
from Features.Text_analysis.text_functions import calculate_word_count

#importing plot functions from Plots.plots
from Features.Plots.plots import plot_top_values
from Features.Plots.plots import plot_word_count_distribution
from Features.Plots.plots import plot_and_save_time_related_data

#importing time functions
from Features.Time_analysis.time_functions import extract_date_time_features

#from Features.Text_analysis.dataset_functions import extract_kmdb_string

#---------------------------------------- Start of the analysis ----------------------------------------#
start_time = time.time()
print(f"Start time -> \n{start_time}\n")

#---------------------------------------- Creation of folder for outputs ----------------------------------------#
# Define the initial folder path
folder_path_input = os.getcwd()
# Input the name of the new folder
new_folder = input("Please input the name of the folder to be created (where the outputs will be sent) -> ")
# Combine the input folder path with the new folder name
outputs_folder = os.path.join(folder_path_input, new_folder)
# Create the preliminary folder
create_folder(outputs_folder)
#invoking the select file feature
selected_file = select_file(folder_path_input)
if selected_file:
    # Process the selected file (e.g., read it into a DataFrame)
    if selected_file.endswith('.xlsx'):
        # Read Excel file
        df = pd.read_excel(selected_file)
    elif selected_file.endswith('.csv'):
        # Read CSV file
        df = pd.read_csv(selected_file)
    print(f"Selected file -> {selected_file}\n")
    # Extract the root of the selected file to create the final folder name
    file_root, file_extension = os.path.splitext(selected_file)
    # Combine the preliminary folder with the file root to create the final output folder
    outputs_folder = os.path.join(outputs_folder, file_root)
    # Create the final output folder
    create_folder(outputs_folder)
else:
    print("No valid file selected.")

# Ensure the output directory exists
#---------------------------------------- Reading the original data and preliminary missing values analysis ----------------------------------------#
df_data = df.copy()
df_data_initial = df_data.copy()

# Calculate the statistics to be sent to the document
filled_stats = calculate_filled_fields(df_data_initial)

# Format the column names to lowercase and with no spaces
df_data_formated = format_column_names(df_data)
#df_data.columns = df_data.columns.str.strip()
#df_data.columns = df_data.columns.str.lower().str.replace(" ", "_").str.strip()
#print(df_data.columns)
df_initial_description = df_data.columns.tolist()

# Handling missing values
missing_values = missing_values_table(df_data)
#Drop the columns having more than 20% missing values
drop_indexes = missing_values[missing_values['% of Total Values'] > 20].index.tolist()

#---------------------------------------- No >=20% missing values in the dataset and filling the missing values with "Not Available" ----------------------------------------#
df_data = df_data.drop(drop_indexes, axis=1)
#print(f"Columns with >20% values -> \n{df_data.columns}\n")
df_data = handle_missing_values(df_data)
df_data_definitive = df_data.columns.tolist()

# time_columns will be used for time functions
print("Select the fields for time analysis")
selected_df, time_columns = select_columns_for_analysis(df_data, df_data.columns)
#df_data[time_columns] = df_data[time_columns].apply(lambda x: pd.to_datetime('1899-12-30') + pd.to_timedelta(x, unit='D'))
#time_columns = ['created', 'updated']

# feature_columns will be used for ploting functions
print("Select the fields for feature analysis")
selected_df, feature_columns = select_columns_for_analysis(df_data, df_data.columns)
#feature_columns = ['priority', 'state', 'assignment_group', 'assigned_to', 'task_type', 'impact', 'urgency', 'category', 'subcategory']

# text_columns will be used for text cleaning functions and word count
print("Select the fields for text analysis")
selected_df, text_columns = select_columns_for_analysis(df_data, df_data.columns)
#to_be_printed_for_text_analysis  = f"The fields of interest for text analysis (cleansing and word count) are: {text_columns}."
#text_columns = ['short_description', 'description', 'resolution_notes', 'resolution_code', 'additional_comments', 'comments_and_work_notes']
# columns to analyze and create the document
columns_to_analyze = feature_columns
#---------------------------------------- Text cleaning of the sub-dataset ----------------------------------------#
for col in text_columns:
  # Create the new column name with "_cleaned" suffix
  cleaned_col_name = f"{col}_cleaned"
  # Modify process_text to return a Series with cleaned text (if needed)
  df_data[cleaned_col_name] = df_data[col].apply(process_text)
  
cleaned_columns = [f"{col}_cleaned" for col in text_columns]
#---------------------------------------- Calculating and plotting the word count ----------------------------------------#
df_data_cleaned = calculate_word_count(df_data, cleaned_columns)
print(df_data_cleaned.columns)
field = "word_count"
plot_word_count_distribution(df_data_cleaned, field, outputs_folder)

#---------------------------------------- Plot of top values for sub datasets ----------------------------------------#
feature_top_values = get_top_values(df_data, columns_to_analyze)
plot_top_values(feature_top_values, outputs_folder)

##---------------------------------------- Plot of time features ----------------------------------------#
##df_data['created_dt'] = pd.to_datetime('1899-12-30') + pd.to_timedelta(df_data['sys_created_on'], unit='D')
##df_data['resolved_dt'] = pd.to_datetime('1899-12-30') + pd.to_timedelta(df_data['resolved_at'], unit='D')
df_data_time = extract_date_time_features(df_data, time_columns)
#df_data_time.to_excel(f"{file_root}_output.xlsx", sheet_name="Date_columns", index=False, engine='openpyxl')
plot_and_save_time_related_data(df_data_time, outputs_folder)

#---------------------------------------- End of the analysis ----------------------------------------#
# Creating the document
document = docx.Document()

# Add the name of the document
document.add_paragraph(f"Document generated on {datetime.today().strftime('%m-%d-%Y')}")
document.add_paragraph()
document.add_heading(f"Data Analysis for {file_root}", level=1)

# Creating the table for initial fields received
# Add a title for the table
document.add_heading("Initial Description of the Data Received", level=2)
# Add a table to the document
table = document.add_table(rows=1, cols=len(df.columns))
document.add_heading("Column Fill Statistics", level=2)
table = document.add_table(rows=1, cols=len(filled_stats.columns))
table.style = 'Table Grid'
# Add headers to the table
header_row = table.rows[0]
for idx, col_name in enumerate(filled_stats.columns):
    header_row.cells[idx].text = col_name

# Add rows with statistics
for _, row in filled_stats.iterrows():
    table_row = table.add_row()
    for idx, value in enumerate(row):
        table_row.cells[idx].text = str(value)
document.add_section()

document.add_paragraph(f"Original data preprocessed as field names to lowercase and no spaces {df_data_formated.columns.tolist()}.")
document.add_paragraph(f"The fields to be drop (>20% missing values) are {drop_indexes}.")
document.add_paragraph(f"Definitive fields to be processed and analyzed are {df_data_definitive}.")

# document adding the unique values per column
document.add_heading("Unique values", level=2)
for column in feature_columns:    
    if column in df_data.columns:
        unique_count = df_data[column].nunique()
        document.add_paragraph(f"Unique {column} are = {unique_count}")
    
document.add_paragraph()
document.add_heading("Notes", level=2)
document.add_paragraph()
##################
plot_path = os.path.join(outputs_folder, "word_count_distribution.png")
document.add_heading("Word count", level=2)
document.add_picture(plot_path, width=Inches(5.5))
#word_count_distribution.png
##################

document.add_page_break()        

for col in columns_to_analyze:
    # Construct file path
    plot_path = os.path.join(outputs_folder, f"{col}_top_values.png")
    document.add_heading(f"Analysis of {col} field.", level=2)

    if os.path.isfile(plot_path):
        # Add the plot to the document
        document.add_picture(plot_path, width=Inches(5.5))
    else:
        # Add a note for missing plots
        document.add_paragraph(f"Plot for {col} is not available at {plot_path}.")
        
    document.add_section()

for col in df_data_time:
    # Construct file path
    plot_path = os.path.join(outputs_folder, f'by_{col}.png')
    if col.endswith("_day_of_week"):
        continue  # Skip columns ending with "_day_of_week"
    
    document.add_heading(f"Analysis of {col} field.", level=2)

    if os.path.isfile(plot_path):
        # Add the plot to the document
        document.add_picture(plot_path, width=Inches(5.5))
    else:
        # Add a note for missing plots
        document.add_paragraph(f"Plot for {col} is not available at {plot_path}.")
        
    document.add_section()

# Save the document
outputs_path = os.path.join(outputs_folder, file_root)
print(f"This is the route: {outputs_path}")

document.save(f"{file_root}.docx")
# Creating the excel workbook
#df_data_time.to_excel(f"{file_root}_output.xlsx", sheet_name = "Date_columns", index=False)
df_data.to_excel(f"{file_root}_output.xlsx", sheet_name = "Output_cleaned", index=False)
end_time = time.time()
duration = end_time - start_time
print(f"Duration (in sec) -> \n{duration}")